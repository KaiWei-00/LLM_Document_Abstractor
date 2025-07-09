import io
from docx import Document
from fastapi import UploadFile


async def extract_text_from_docx(file: UploadFile) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        file: DOCX file uploaded by the user
    
    Returns:
        str: Extracted text content
    
    Raises:
        Exception: If there's an error during DOCX processing
    """
    try:
        # Read the file content
        content = await file.read()
        
        # Use python-docx to extract text
        doc = Document(io.BytesIO(content))
        
        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs]
        text = "\n".join(paragraphs)
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Clean up the text
        text = text.strip()
        
        # Reset file pointer for potential later use
        await file.seek(0)
        
        return text
    except Exception as e:
        # Re-raise with more context
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
